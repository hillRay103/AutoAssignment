#!/usr/bin/env python3
"""
AutoAssignment v6 – Document Generator
Requires: pip install python-docx reportlab
PDF uses reportlab directly — no docx2pdf, no LibreOffice needed.
"""
import sys, json
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# ── Colour constants ──────────────────────────────────────────────────────────
def rgb(h): h=h.lstrip('#'); return RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16))
NAVY="1A3A6C"; BLUE="1F6FB2"; DGRAY="1E1E1E"; MGRAY="777777"; WHITE="FFFFFF"; GREEN="145A32"
CODE_BG={"lightgray":"F4F4F4","lightblue":"EBF5FB","darkgray":"1E1E1E","white":"FFFFFF"}
CODE_FG={"lightgray":DGRAY,"lightblue":NAVY,"darkgray":"D4D4D4","white":DGRAY}
OUT_BG ={"lightgray":"EAF7EE","lightblue":"EAF7EE","darkgray":"051A0F","white":"EAF7EE"}

# ── XML helpers ───────────────────────────────────────────────────────────────
def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill.lstrip('#')); tcPr.append(shd)

def bdr(cell, color="AAAAAA", sz=4):
    tcPr=cell._tc.get_or_add_tcPr(); tcB=OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        el=OxmlElement(f'w:{e}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),str(sz))
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),color.lstrip('#')); tcB.append(el)
    tcPr.append(tcB)

def nobdr(cell):
    tcPr=cell._tc.get_or_add_tcPr(); tcB=OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        el=OxmlElement(f'w:{e}'); el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),'FFFFFF'); tcB.append(el)
    tcPr.append(tcB)

def cellpad(cell, t=50, b=50, l=80, r=80):
    tcPr=cell._tc.get_or_add_tcPr(); mar=OxmlElement('w:tcMar')
    for side,val in [('top',t),('left',l),('bottom',b),('right',r)]:
        el=OxmlElement(f'w:{side}'); el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa'); mar.append(el)
    tcPr.append(mar)

def setwid(cell, inches):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=OxmlElement('w:tcW')
    tcW.set(qn('w:w'),str(int(inches*1440))); tcW.set(qn('w:type'),'dxa'); tcPr.append(tcW)

def addp(doc, text, bold=False, italic=False, size=11, color=DGRAY,
         align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4, font='Calibri'):
    p=doc.add_paragraph(); p.alignment=align
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.italic=italic
    r.font.name=font; r.font.size=Pt(size); r.font.color.rgb=rgb(color)
    return p

def divline(doc, color=NAVY):
    dp=doc.add_paragraph(); dp.paragraph_format.space_before=Pt(1); dp.paragraph_format.space_after=Pt(1)
    pPr=dp._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),color); pBdr.append(bot); pPr.append(pBdr)

def footer_pgnum(section, name, title, font):
    footer=section.footer
    fp=footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run(f"{name}  |  {title}  |  Page ")
    r.font.size=Pt(9); r.font.color.rgb=rgb(MGRAY); r.font.name=font
    for stage in ['begin', 'instrText', 'end']:
        el=OxmlElement('w:r'); rp=OxmlElement('w:rPr')
        sz=OxmlElement('w:sz'); sz.set(qn('w:val'),'18'); rp.append(sz)
        cl=OxmlElement('w:color'); cl.set(qn('w:val'),MGRAY); rp.append(cl)
        el.append(rp)
        if stage=='begin':
            fc=OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'),'begin'); el.append(fc)
        elif stage=='instrText':
            it=OxmlElement('w:instrText'); it.text='PAGE'; el.append(it)
        else:
            fc=OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'),'end'); el.append(fc)
        fp._p.append(el)

def code_block(doc, label, code_text, output_text, theme, fs, font):
    cbg=CODE_BG.get(theme,'F4F4F4'); cfg=CODE_FG.get(theme,DGRAY); obg=OUT_BG.get(theme,'EAF7EE')
    # Section heading
    lp=doc.add_paragraph(); lp.paragraph_format.space_before=Pt(10); lp.paragraph_format.space_after=Pt(4)
    lr=lp.add_run(f"▸ {label}"); lr.bold=True; lr.font.size=Pt(13); lr.font.color.rgb=rgb(NAVY); lr.font.name=font

    def block(header, content, bg, fg, hcol, bcol):
        tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
        c=tbl.cell(0,0); shade(c,bg); bdr(c,color=bcol,sz=6); cellpad(c,60,60,100,100)
        hp=c.paragraphs[0]; hp.paragraph_format.space_before=Pt(4); hp.paragraph_format.space_after=Pt(3)
        hr=hp.add_run(header); hr.bold=True; hr.font.size=Pt(fs); hr.font.color.rgb=rgb(hcol); hr.font.name=font
        for line in (content or '').splitlines():
            lp2=c.add_paragraph(line if line.strip() else ' ')
            lp2.paragraph_format.space_before=Pt(0); lp2.paragraph_format.space_after=Pt(0)
            lr2=lp2.runs[0] if lp2.runs else lp2.add_run(line)
            lr2.font.name='Courier New'; lr2.font.size=Pt(fs); lr2.font.color.rgb=rgb(fg)
        doc.add_paragraph().paragraph_format.space_after=Pt(3)

    block("Source Code:", code_text, cbg, cfg, BLUE, BLUE)
    if output_text and output_text.strip():
        block("Program Output:", output_text, obg, DGRAY, GREEN, GREEN)


# ── Build DOCX ─────────────────────────────────────────────────────────────────
def build_docx(data):
    doc=Document(); sec=doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.left_margin=Inches(1); sec.right_margin=Inches(1)
    sec.top_margin=Inches(1); sec.bottom_margin=Inches(1)
    font=data.get('fontChoice','Calibri')
    doc.styles['Normal'].font.name=font; doc.styles['Normal'].font.size=Pt(11)

    tpl=data.get('template','B')
    name=data.get('studentName',''); reg=data.get('regNumber','')
    course=data.get('courseName',''); course_c=data.get('courseCode','')
    section=data.get('section',''); instructor=data.get('submittedTo','')
    dept=data.get('department',''); uni=data.get('university','')
    prog=data.get('programName',''); sem=data.get('semester','')
    batch=data.get('batch',''); credits=data.get('creditHours','')
    assign_no=data.get('assignmentNumber',1); assign_ttl=data.get('assignmentTitle','')
    sub_date=data.get('submissionDate',datetime.now().strftime('%d-%m-%Y'))
    clo=data.get('cloLabel',''); project=data.get('projectTitle','')
    questions=data.get('questions',[]); theme=data.get('codeTheme','lightgray')
    fs=int(data.get('codeFontSize','11'))
    full_title=f"Assignment No. {assign_no}"+(f" — {assign_ttl}" if assign_ttl else "")
    course_str=course+(f" ({course_c})" if course_c else "")
    footer_pgnum(sec, name, full_title, font)

    # ══ TEMPLATE A ═══════════════════════════════════════════════════════════
    if tpl=='A':
        if uni:
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(1)
            r=p.add_run(uni); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=rgb(NAVY); r.font.name=font
        if dept:
            p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(1)
            r2=p2.add_run(dept); r2.bold=True; r2.font.size=Pt(11); r2.font.color.rgb=rgb(NAVY); r2.font.name=font
        divline(doc,NAVY); divline(doc,NAVY)
        doc.add_paragraph().paragraph_format.space_after=Pt(5)
        specs=[
            [("Course Title:",True,1.3),(course,False,1.8),("Course Code:",True,1.0),(course_c,False,0.7),("Credit Hours:",True,1.0),(credits,False,0.7)],
            [("Instructor:",True,1.3),(instructor,False,1.8),("Program:",True,1.0),(prog or "BS(CS)",False,2.4)],
            [("Semester:",True,1.3),(sem,False,0.6),("Batch:",True,0.8),(batch,False,0.7),("Section:",True,0.8),(section,False,2.3)],
            [("Date:",True,1.3),(sub_date,False,1.8),("",True,1.0),("",False,3.4)],
            [("Name:",True,1.3),(name,False,1.8),("Reg No:",True,1.0),(reg,False,3.4)],
        ]
        for row_spec in specs:
            tbl=doc.add_table(rows=1,cols=len(row_spec)); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
            for ci,(lbl,is_lbl,w) in enumerate(row_spec):
                c=tbl.cell(0,ci); shade(c,'EBEBEB' if is_lbl else WHITE); bdr(c,'888888',sz=4)
                cellpad(c,40,40,80,80); setwid(c,w)
                cp=c.paragraphs[0]; cp.paragraph_format.space_before=Pt(3); cp.paragraph_format.space_after=Pt(3)
                r=cp.add_run(lbl); r.bold=is_lbl; r.font.size=Pt(10); r.font.name=font
                r.font.color.rgb=rgb(NAVY if is_lbl else DGRAY)
        doc.add_paragraph().paragraph_format.space_after=Pt(10)

    # ══ TEMPLATE B ═══════════════════════════════════════════════════════════
    elif tpl=='B':
        tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER; tp.paragraph_format.space_after=Pt(8)
        tr=tp.add_run(full_title); tr.bold=True; tr.font.size=Pt(15); tr.font.color.rgb=rgb(NAVY); tr.font.name=font
        rows=[[("Student Name:",True),(name,False),("Reg No.:",True),(reg,False)],
              [("Course:",True),(course_str,False),("Section:",True),(section,False)],
              [("Instructor:",True),(instructor,False),("Date:",True),(sub_date,False)]]
        if sem: rows.append([("Semester:",True),(sem,False),("Batch:",True),(batch,False)])
        widths=[1.5,2.1,1.4,1.8]
        htbl=doc.add_table(rows=len(rows),cols=4); htbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        for ri,row in enumerate(rows):
            for ci,(txt,is_lbl) in enumerate(row):
                c=htbl.cell(ri,ci); shade(c,'EBF0F7' if is_lbl else WHITE); bdr(c,'999999',sz=4)
                cellpad(c,50,50,80,80); setwid(c,widths[ci])
                cp=c.paragraphs[0]; cp.paragraph_format.space_before=Pt(3); cp.paragraph_format.space_after=Pt(3)
                r=cp.add_run(txt); r.bold=is_lbl; r.font.size=Pt(10); r.font.name=font
                r.font.color.rgb=rgb(NAVY if is_lbl else DGRAY)
        doc.add_paragraph()

    # ══ TEMPLATE C ═══════════════════════════════════════════════════════════
    elif tpl=='C':
        addp(doc,full_title,bold=True,size=15,color=NAVY,align=WD_ALIGN_PARAGRAPH.CENTER,sb=0,sa=3,font=font)
        addp(doc,f"{name}  |  {reg}",size=11,color=DGRAY,align=WD_ALIGN_PARAGRAPH.CENTER,sb=0,sa=2,font=font)
        parts=[course_str,f"Section {section}",instructor]
        addp(doc,'  |  '.join(p for p in parts if p),size=10,color=MGRAY,align=WD_ALIGN_PARAGRAPH.CENTER,sb=0,sa=2,font=font)
        addp(doc,f"Date: {sub_date}",size=10,color=MGRAY,align=WD_ALIGN_PARAGRAPH.CENTER,sb=0,sa=6,font=font)
        divline(doc,BLUE); doc.add_paragraph().paragraph_format.space_after=Pt(4)

    # ══ TEMPLATE D — Cover Page ═══════════════════════════════════════════════
    elif tpl=='D':
        cv=doc.add_table(rows=1,cols=2); cv.alignment=WD_TABLE_ALIGNMENT.LEFT
        lc=cv.cell(0,0); rc=cv.cell(0,1)
        shade(lc,'1A3A6C'); setwid(lc,0.25); nobdr(lc); nobdr(rc)
        setwid(rc,6.0); cellpad(rc,40,40,160,40)
        lc.paragraphs[0].paragraph_format.space_before=Pt(0)
        lc.paragraphs[0].paragraph_format.space_after=Pt(0)
        # CLO tag
        rp0=rc.paragraphs[0]; rp0.paragraph_format.space_before=Pt(8); rp0.paragraph_format.space_after=Pt(4)
        if clo:
            tag_t=rc.add_table(rows=1,cols=1); tc=tag_t.cell(0,0)
            shade(tc,'1A6B8A'); cellpad(tc,30,30,80,80); setwid(tc,1.5); nobdr(tc)
            tp2=tc.paragraphs[0]; tp2.alignment=WD_ALIGN_PARAGRAPH.LEFT
            tr2=tp2.add_run(clo); tr2.bold=True; tr2.font.size=Pt(10); tr2.font.color.rgb=rgb('FFFFFF'); tr2.font.name=font
            rc.add_paragraph().paragraph_format.space_after=Pt(6)
        # Course title
        ct=rc.add_paragraph(); ct.paragraph_format.space_after=Pt(4)
        cr_=ct.add_run(course); cr_.bold=True; cr_.font.size=Pt(20); cr_.font.color.rgb=rgb(DGRAY); cr_.font.name=font
        if assign_ttl:
            d2=rc.add_paragraph(); d2.paragraph_format.space_after=Pt(10)
            dr=d2.add_run(assign_ttl); dr.font.size=Pt(11); dr.font.color.rgb=rgb(MGRAY); dr.font.name=font
        # Divider
        div_p=rc.add_paragraph(); div_p.paragraph_format.space_before=Pt(6); div_p.paragraph_format.space_after=Pt(6)
        dPr=div_p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
        bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),NAVY); pBdr.append(bot); dPr.append(pBdr)
        # Project title
        if project:
            pt_p=rc.add_paragraph(); pt_p.paragraph_format.space_after=Pt(4)
            ptr=pt_p.add_run("Project Title:"); ptr.bold=True; ptr.font.size=Pt(11); ptr.font.color.rgb=rgb(NAVY); ptr.font.name=font
            pt_p2=rc.add_paragraph(); pt_p2.paragraph_format.space_after=Pt(14)
            ptr2=pt_p2.add_run(project); ptr2.bold=True; ptr2.italic=True
            ptr2.font.size=Pt(13); ptr2.font.color.rgb=rgb(NAVY); ptr2.font.name=font
            div2=rc.add_paragraph(); div2.paragraph_format.space_before=Pt(2); div2.paragraph_format.space_after=Pt(12)
            d2Pr=div2._p.get_or_add_pPr(); pBdr2=OxmlElement('w:pBdr')
            bot2=OxmlElement('w:bottom'); bot2.set(qn('w:val'),'single'); bot2.set(qn('w:sz'),'4')
            bot2.set(qn('w:space'),'1'); bot2.set(qn('w:color'),NAVY); pBdr2.append(bot2); d2Pr.append(pBdr2)
        # Submission footer
        for txt, bold_, col in [
            (f"Submitted to: {instructor}", True,  DGRAY),
            ("SUBMITTED BY:",              True,  DGRAY),
            (f"{reg}    {name.upper()}",   True,  BLUE),
            (f"DATE OF SUBMISSION:\n{sub_date}", False, DGRAY),
        ]:
            sp=rc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; sp.paragraph_format.space_after=Pt(2)
            sr=sp.add_run(txt); sr.bold=bold_; sr.font.size=Pt(10); sr.font.color.rgb=rgb(col); sr.font.name=font
        doc.add_page_break()

    # ══ QUESTIONS ════════════════════════════════════════════════════════════
    for i,q in enumerate(questions):
        label=q.get('label',f'Question {i+1}'); qtext=q.get('text','')
        code=q.get('code',''); output=q.get('output',''); notes=q.get('notes','')
        if qtext: addp(doc,f"Q: {qtext}",italic=True,size=10,color=MGRAY,sb=4,sa=2,font=font)
        code_block(doc,label,code,output,theme,fs,font)
        if notes: addp(doc,f"Note: {notes}",italic=True,size=10,color=MGRAY,sb=0,sa=8,font=font)
    return doc


# ── Build PDF using ReportLab ─────────────────────────────────────────────────
def build_pdf(data, out_path):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    except ImportError:
        print("reportlab not installed. Run: pip install reportlab", file=sys.stderr)
        sys.exit(1)

    font       = data.get('fontChoice', 'Calibri')
    # ReportLab uses Helvetica for Calibri-like, Times-Roman for Times New Roman
    rl_font    = 'Times-Roman' if 'Times' in font else 'Helvetica'
    rl_bold    = 'Times-Bold'  if 'Times' in font else 'Helvetica-Bold'
    rl_italic  = 'Times-Italic' if 'Times' in font else 'Helvetica-Oblique'

    name       = data.get('studentName',''); reg=data.get('regNumber','')
    course     = data.get('courseName',''); course_c=data.get('courseCode','')
    section    = data.get('section',''); instructor=data.get('submittedTo','')
    dept       = data.get('department',''); uni=data.get('university','')
    sem        = data.get('semester',''); batch=data.get('batch',''); credits=data.get('creditHours','')
    assign_no  = data.get('assignmentNumber',1); assign_ttl=data.get('assignmentTitle','')
    sub_date   = data.get('submissionDate',datetime.now().strftime('%d-%m-%Y'))
    questions  = data.get('questions',[])
    tpl        = data.get('template','B')
    course_str = course+(f" ({course_c})" if course_c else "")
    full_title = f"Assignment No. {assign_no}"+(f" — {assign_ttl}" if assign_ttl else "")

    NAVY_RL  = colors.HexColor('#1A3A6C')
    BLUE_RL  = colors.HexColor('#1F6FB2')
    GREEN_RL = colors.HexColor('#145A32')
    GRAY_RL  = colors.HexColor('#777777')
    LGRAY_RL = colors.HexColor('#F4F4F4')
    LBLUE_RL = colors.HexColor('#EBF5FB')
    OGRN_RL  = colors.HexColor('#EAF7EE')

    doc = SimpleDocTemplate(out_path, pagesize=letter,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)

    styles = getSampleStyleSheet()
    def sty(name2, **kw):
        defaults = dict(fontName=rl_font, fontSize=11, textColor=colors.HexColor('#1E1E1E'))
        defaults.update(kw)
        return ParagraphStyle(name2, **defaults)

    story = []

    # ── Header based on template ──────────────────────────────────────────────
    if tpl == 'A':
        if uni:
            story.append(Paragraph(f'<b>{uni}</b>', sty('uni', fontSize=14, textColor=NAVY_RL, alignment=TA_CENTER)))
            story.append(Spacer(1,4))
        if dept:
            story.append(Paragraph(f'<b>{dept}</b>', sty('dept', fontSize=11, textColor=NAVY_RL, alignment=TA_CENTER)))
        story.append(HRFlowable(width='100%', thickness=2, color=NAVY_RL, spaceAfter=4))
        story.append(HRFlowable(width='100%', thickness=1, color=NAVY_RL, spaceAfter=6))
        tbl_data = [
            ['Course Title:', course, 'Course Code:', course_c, 'Credit Hours:', credits],
            ['Instructor:', instructor, 'Program:', data.get('programName','BS(CS)'), '', ''],
            ['Semester:', sem, 'Batch:', batch, 'Section:', section],
            ['Date:', sub_date, '', '', '', ''],
            ['Name:', name, 'Reg No:', reg, '', ''],
        ]
        col_w = [1.1*inch, 1.4*inch, 1.0*inch, 0.8*inch, 1.0*inch, 0.7*inch]
        t = Table(tbl_data, colWidths=col_w)
        ts = TableStyle([
            ('FONTNAME',    (0,0),(-1,-1), rl_font),
            ('FONTSIZE',    (0,0),(-1,-1), 9),
            ('FONTNAME',    (0,0),(0,-1),  rl_bold),
            ('FONTNAME',    (2,0),(2,-1),  rl_bold),
            ('FONTNAME',    (4,0),(4,-1),  rl_bold),
            ('TEXTCOLOR',   (0,0),(0,-1),  NAVY_RL),
            ('TEXTCOLOR',   (2,0),(2,-1),  NAVY_RL),
            ('TEXTCOLOR',   (4,0),(4,-1),  NAVY_RL),
            ('BACKGROUND',  (0,0),(0,-1),  LGRAY_RL),
            ('BACKGROUND',  (2,0),(2,-1),  LGRAY_RL),
            ('BACKGROUND',  (4,0),(4,-1),  LGRAY_RL),
            ('GRID',        (0,0),(-1,-1), 0.5, colors.HexColor('#888888')),
            ('TOPPADDING',  (0,0),(-1,-1), 3),
            ('BOTTOMPADDING',(0,0),(-1,-1),3),
        ])
        t.setStyle(ts)
        story.append(t); story.append(Spacer(1,12))

    elif tpl == 'B':
        story.append(Paragraph(f'<b>{full_title}</b>', sty('ft', fontSize=15, textColor=NAVY_RL, alignment=TA_CENTER)))
        story.append(Spacer(1,8))
        tbl_data = [
            ['Student Name:', name, 'Reg No.:', reg],
            ['Course:', course_str, 'Section:', section],
            ['Instructor:', instructor, 'Date:', sub_date],
        ]
        if sem: tbl_data.append(['Semester:', sem, 'Batch:', batch])
        col_w = [1.3*inch, 1.9*inch, 1.2*inch, 1.6*inch]
        t = Table(tbl_data, colWidths=col_w)
        ts = TableStyle([
            ('FONTNAME',    (0,0),(-1,-1), rl_font), ('FONTSIZE',(0,0),(-1,-1),9),
            ('FONTNAME',    (0,0),(0,-1),  rl_bold), ('FONTNAME',(2,0),(2,-1),rl_bold),
            ('TEXTCOLOR',   (0,0),(0,-1),  NAVY_RL), ('TEXTCOLOR',(2,0),(2,-1),NAVY_RL),
            ('BACKGROUND',  (0,0),(0,-1),  LBLUE_RL),('BACKGROUND',(2,0),(2,-1),LBLUE_RL),
            ('GRID',        (0,0),(-1,-1), 0.5, colors.HexColor('#999999')),
            ('TOPPADDING',  (0,0),(-1,-1), 4), ('BOTTOMPADDING',(0,0),(-1,-1),4),
        ])
        t.setStyle(ts); story.append(t); story.append(Spacer(1,10))

    elif tpl == 'C':
        story.append(Paragraph(f'<b>{full_title}</b>', sty('ft2', fontSize=15, textColor=NAVY_RL, alignment=TA_CENTER)))
        story.append(Paragraph(f'{name}  |  {reg}', sty('sub', fontSize=11, alignment=TA_CENTER)))
        story.append(Paragraph(f'{course_str}  |  Section {section}  |  {instructor}', sty('sub2', fontSize=10, textColor=GRAY_RL, alignment=TA_CENTER)))
        story.append(HRFlowable(width='100%', thickness=1.5, color=BLUE_RL, spaceBefore=4, spaceAfter=8))

    else:  # Template D — Cover Page
        clo        = data.get('cloLabel', '')
        project    = data.get('projectTitle', '')
        prog_pdf   = data.get('programName', '')
        # Side-bar + body as a two-column table
        cover_tbl = Table([['', '']], colWidths=[0.25*inch, 6.0*inch])
        cover_tbl.setStyle(TableStyle([
            ('BACKGROUND',   (0,0),(0,0), NAVY_RL),
            ('VALIGN',       (0,0),(-1,-1), 'TOP'),
            ('TOPPADDING',   (0,0),(-1,-1), 0),
            ('BOTTOMPADDING',(0,0),(-1,-1), 0),
            ('LEFTPADDING',  (0,0),(-1,-1), 0),
            ('RIGHTPADDING', (0,0),(-1,-1), 0),
            ('LINEAFTER',    (0,0),(0,-1), 0, colors.white),
        ]))
        story.append(cover_tbl)
        story.append(Spacer(1, 6))
        if clo:
            clo_tbl = Table([[clo]], colWidths=[1.5*inch])
            clo_tbl.setStyle(TableStyle([
                ('BACKGROUND',    (0,0),(-1,-1), colors.HexColor('#1A6B8A')),
                ('TEXTCOLOR',     (0,0),(-1,-1), colors.white),
                ('FONTNAME',      (0,0),(-1,-1), rl_bold),
                ('FONTSIZE',      (0,0),(-1,-1), 10),
                ('TOPPADDING',    (0,0),(-1,-1), 3),
                ('BOTTOMPADDING', (0,0),(-1,-1), 3),
                ('LEFTPADDING',   (0,0),(-1,-1), 8),
            ]))
            story.append(clo_tbl)
            story.append(Spacer(1, 6))
        story.append(Paragraph(f'<b>{course}</b>', sty('cv_course', fontSize=20, textColor=colors.HexColor('#1E1E1E'))))
        if assign_ttl:
            story.append(Spacer(1, 4))
            story.append(Paragraph(assign_ttl, sty('cv_sub', fontSize=11, textColor=GRAY_RL)))
        story.append(HRFlowable(width='100%', thickness=1.5, color=NAVY_RL, spaceBefore=6, spaceAfter=6))
        if project:
            story.append(Paragraph('<b>Project Title:</b>', sty('cv_ptlbl', fontSize=11, textColor=NAVY_RL)))
            story.append(Paragraph(f'<i><b>{project}</b></i>', sty('cv_pt', fontSize=13, textColor=NAVY_RL)))
            story.append(HRFlowable(width='100%', thickness=1, color=NAVY_RL, spaceBefore=4, spaceAfter=10))
        story.append(Paragraph(f'<b>Submitted to: {instructor}</b>', sty('cv_f1', fontSize=10, alignment=TA_RIGHT)))
        story.append(Paragraph('<b>SUBMITTED BY:</b>', sty('cv_f2', fontSize=10, alignment=TA_RIGHT)))
        story.append(Paragraph(f'<b><font color="#1F6FB2">{reg}    {name.upper()}</font></b>', sty('cv_f3', fontSize=10, alignment=TA_RIGHT)))
        story.append(Paragraph(f'DATE OF SUBMISSION: {sub_date}', sty('cv_f4', fontSize=10, alignment=TA_RIGHT)))
        from reportlab.platypus import PageBreak
        story.append(PageBreak())

    # ── Questions ─────────────────────────────────────────────────────────────
    code_font = 'Courier'
    for q in questions:
        label  = q.get('label','Question')
        qtext  = q.get('text','')
        code   = q.get('code','')
        output = q.get('output','')
        notes  = q.get('notes','')

        story.append(Spacer(1,8))
        story.append(Paragraph(f'<b>▸ {label}</b>', sty('qlbl', fontSize=13, textColor=NAVY_RL)))
        story.append(Spacer(1,4))

        if qtext:
            story.append(Paragraph(f'<i>Q: {qtext}</i>', sty('qt', fontSize=10, textColor=GRAY_RL)))
            story.append(Spacer(1,3))

        if code:
            code_lines = [['Source Code:']] + [[line] for line in (code).splitlines()]
            ct = Table(code_lines, colWidths=[6*inch])
            ct.setStyle(TableStyle([
                ('FONTNAME',    (0,0),(0,0), rl_bold),    ('FONTSIZE',(0,0),(0,0),10),
                ('TEXTCOLOR',   (0,0),(0,0), BLUE_RL),
                ('FONTNAME',    (0,1),(-1,-1), code_font),('FONTSIZE',(0,1),(-1,-1),9),
                ('TEXTCOLOR',   (0,1),(-1,-1), colors.HexColor('#1E1E1E')),
                ('BACKGROUND',  (0,0),(-1,-1), LGRAY_RL),
                ('BOX',         (0,0),(-1,-1), 1, BLUE_RL),
                ('TOPPADDING',  (0,0),(-1,-1), 2), ('BOTTOMPADDING',(0,0),(-1,-1),2),
                ('LEFTPADDING', (0,0),(-1,-1), 6), ('RIGHTPADDING',(0,0),(-1,-1),6),
            ]))
            story.append(ct); story.append(Spacer(1,4))

        if output and output.strip():
            out_lines = [['Program Output:']] + [[line] for line in output.splitlines()]
            ot = Table(out_lines, colWidths=[6*inch])
            ot.setStyle(TableStyle([
                ('FONTNAME',    (0,0),(0,0), rl_bold),    ('FONTSIZE',(0,0),(0,0),10),
                ('TEXTCOLOR',   (0,0),(0,0), GREEN_RL),
                ('FONTNAME',    (0,1),(-1,-1), code_font),('FONTSIZE',(0,1),(-1,-1),9),
                ('TEXTCOLOR',   (0,1),(-1,-1), colors.HexColor('#1E1E1E')),
                ('BACKGROUND',  (0,0),(-1,-1), OGRN_RL),
                ('BOX',         (0,0),(-1,-1), 1, GREEN_RL),
                ('TOPPADDING',  (0,0),(-1,-1), 2), ('BOTTOMPADDING',(0,0),(-1,-1),2),
                ('LEFTPADDING', (0,0),(-1,-1), 6), ('RIGHTPADDING',(0,0),(-1,-1),6),
            ]))
            story.append(ot); story.append(Spacer(1,4))

        if notes:
            story.append(Paragraph(f'<i>Note: {notes}</i>', sty('nt', fontSize=10, textColor=GRAY_RL)))

    # Footer
    def footer_cb(canvas, doc2):
        canvas.saveState()
        canvas.setFont(rl_font, 8)
        canvas.setFillColor(GRAY_RL)
        txt = f"{name}  |  {full_title}  |  Page {doc2.page}"
        canvas.drawCentredString(letter[0]/2, 0.5*inch, txt)
        canvas.restoreState()

    doc.build(story, onFirstPage=footer_cb, onLaterPages=footer_cb)


# ── Entry point ───────────────────────────────────────────────────────────────
def main():
    try:
        raw  = sys.stdin.buffer.read().decode('utf-8')
        data = json.loads(raw)
    except Exception as ex:
        print(f"Input error: {ex}", file=sys.stderr); sys.exit(1)

    out_path = data.get('outputPath', 'Assignment.docx')
    fmt      = data.get('outputFormat', 'docx')

    try:
        if fmt == 'pdf':
            build_pdf(data, out_path)
        else:
            doc = build_docx(data)
            doc.save(out_path)
        print(f"OK:{out_path}")
    except Exception as ex:
        print(f"Generator error: {ex}", file=sys.stderr); sys.exit(1)

if __name__ == '__main__':
    main()
